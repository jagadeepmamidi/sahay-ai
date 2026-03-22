"""
Document Processor
==================

Processes uploaded documents (PDF, DOCX, TXT) and extracts text for ingestion
into the vector database.

Author: Jagadeep Mamidi
"""

import os
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import hashlib

from app.core.config import get_settings
from app.rag.hybrid_retriever import get_retriever

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """
    Processes various document formats and prepares content for ingestion.
    
    Supports:
    - PDF files
    - Plain text files
    - Markdown files
    - Word documents (DOCX)
    """
    
    def __init__(self):
        self.retriever = get_retriever()
        self.chunk_size = 1000
        self.chunk_overlap = 150
    
    async def process_document(
        self, 
        file_path: str, 
        filename: str,
        metadata: Optional[Dict] = None
    ) -> Dict[str, Any]:
        """
        Process a document file and add to vector database.
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            metadata: Optional metadata to attach
            
        Returns:
            Dict with success status and processing details
        """
        try:
            # Determine file type
            ext = Path(filename).suffix.lower()
            
            # Extract text based on file type
            if ext == ".pdf":
                text = self._extract_pdf(file_path)
            elif ext == ".docx":
                text = self._extract_docx(file_path)
            elif ext in [".txt", ".md"]:
                text = self._extract_text(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {ext}"
                }
            
            if not text or len(text.strip()) < 50:
                return {
                    "success": False,
                    "error": "Insufficient text extracted from document"
                }
            
            # Create chunks
            chunks = self._chunk_text(text)
            
            # Generate document ID
            doc_hash = hashlib.md5(text.encode()).hexdigest()[:8]
            base_id = f"{Path(filename).stem}_{doc_hash}"
            
            # Prepare metadata
            doc_metadata = {
                "source_file": filename,
                "processed_at": datetime.utcnow().isoformat(),
                "chunk_count": len(chunks),
                **(metadata or {})
            }
            
            # Add chunks to retriever
            for i, chunk in enumerate(chunks):
                chunk_id = f"{base_id}_chunk_{i}"
                chunk_metadata = {
                    **doc_metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                self.retriever.add_document(chunk_id, chunk, chunk_metadata)
            
            logger.info(f"Processed document: {filename} -> {len(chunks)} chunks")
            
            return {
                "success": True,
                "document_id": base_id,
                "chunks_created": len(chunks),
                "total_characters": len(text)
            }
            
        except Exception as e:
            logger.error(f"Document processing error: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import fitz  # PyMuPDF
            
            text_parts = []
            with fitz.open(file_path) as pdf:
                for page_num, page in enumerate(pdf):
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            # Fallback to pypdf
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {i + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            # If python-docx not installed, read as text
            return self._extract_text(file_path)
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text or markdown file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap for better retrieval.
        Uses sentence-aware splitting to avoid breaking mid-sentence.
        """
        # Simple sentence-aware chunking
        sentences = []
        current = ""
        
        for char in text:
            current += char
            if char in ".!?" and len(current) > 50:
                sentences.append(current.strip())
                current = ""
        
        if current.strip():
            sentences.append(current.strip())
        
        # Combine sentences into chunks
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < self.chunk_size:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                if chunks and self.chunk_overlap > 0:
                    # Take end of previous chunk for overlap
                    overlap = chunks[-1][-self.chunk_overlap:]
                    current_chunk = overlap + " " + sentence
                else:
                    current_chunk = sentence
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    async def process_scheme_data(
        self,
        scheme_id: str,
        scheme_name: str,
        category: str,
        content: str,
        benefit_summary: str = "",
        eligibility_summary: str = ""
    ) -> Dict[str, Any]:
        """
        Process structured scheme data and add to vector database.
        
        This is for manually adding scheme information.
        """
        try:
            metadata = {
                "scheme_id": scheme_id,
                "scheme_name": scheme_name,
                "category": category,
                "benefit_summary": benefit_summary,
                "eligibility_summary": eligibility_summary,
                "processed_at": datetime.utcnow().isoformat()
            }
            
            # Chunk the content
            chunks = self._chunk_text(content)
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{scheme_id}_chunk_{i}"
                chunk_metadata = {
                    **metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                self.retriever.add_document(chunk_id, chunk, chunk_metadata)
            
            logger.info(f"Added scheme: {scheme_name} -> {len(chunks)} chunks")
            
            return {
                "success": True,
                "scheme_id": scheme_id,
                "chunks_created": len(chunks)
            }
            
        except Exception as e:
            logger.error(f"Scheme processing error: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }
